from .models import *
# Import docx NOT python-docx 
import docx 
from docx.shared import Pt 
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT,WD_PARAGRAPH_ALIGNMENT 
import jdatetime

# Create an instance of a word document 
def docx_generator(allasnad , shomare_sanad , name_sherkat , date_sanad , category_obj , user1):
    print('1')
    doc = docx.Document() 
    style = doc.styles.add_style('rtl', WD_STYLE_TYPE.PARAGRAPH)
    print('1')

    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    print('1')


    # Add a Title to the document 
    d1=doc.add_heading('حسابداری سند', 0 ) 
    d1.alignment = 1
    ##############################################################################
    table1 = doc.add_table(rows=1, cols=4) 
    print('1')
    


    # Adding heading in the 1st row of the table 
    row = table1.rows[0].cells 
    row[3].text = 'سند تاریخ'
    row[2].text = 'سند شماره'
    row[1].text = 'شرکت اسم'

    print('1')
    
    # Adding data from the list to the table 
    text_finaler1 =''
    name=name_sherkat.split(' ')
    for i in reversed(name):
        text_finaler1+=f'{i} '

    final_date=jdatetime.date.fromgregorian(year=int(date_sanad.strftime("%Y")) , month=int(date_sanad.strftime("%m")) , day=int(date_sanad.strftime("%d")))
    print(final_date)

    row = table1.add_row().cells 
    row[3].text = f"{str(final_date)}"
    row[2].text = f'{shomare_sanad}' 
    row[1].text = f"{text_finaler1}"
    print('1')

    
    # Adding style to a table 
    table1.style = 'Colorful List'


    ###############################################################################
    doc.add_paragraph(style='rtl').add_run("")

    # Table data in a form of list 
    data = ( 
        (1, 'من پسر خوبی هستم'), 
        (2, 'Geek 2'), 
        (3, 'Geek 3') 
    ) 
    print('1')
    
    # Creating a table object 
    table = doc.add_table(rows=1, cols=4) 
    
    # Adding heading in the 1st row of the table 
    row = table.rows[0].cells 
    row[3].text = 'حساب'
    row[2].text = 'شرح'
    row[1].text = 'بدهکار'
    row[0].text = 'بستانکار'
    bedehkar_total = 0
    bestankar_total = 0
    print('1')

    for i in allasnad:
        if i.verified == True:
            obj_sanad=Asnad.objects.get( kol=i.kol , moein=i.moein , tafs = i.tafs , sharhe_hesab = i.sharhe_hesab , bedehkar = i.bedehkar , bestankar =i.bestankar , user=user1,category=category_obj , verified=True)
            row = table.add_row().cells 
            name = obj_sanad.sharhe_hesab.split(' ')
            text_finaler =''
            print(obj_sanad.sharhe_hesab)
            for i in reversed(name):
                text_finaler+=f'{i} '
            bedehkar_total += int(obj_sanad.bedehkar)
            bestankar_total += int(obj_sanad.bestankar)

            obj_sanad.bedehkar = "{:,}".format(obj_sanad.bedehkar)
            obj_sanad.bestankar = "{:,}".format(obj_sanad.bestankar)
            
            ########### name asnad ####################
            name_sanad_kol = obj_sanad.kol.split(' ')
            name_sanad_moein = obj_sanad.moein.split(' ')
            name_sanad_tafs = obj_sanad.tafs.split(' ')

            final_name_sanad = ''

            for i in reversed(name_sanad_kol):
                final_name_sanad+=f'{i}'
            final_name_sanad +=f'-'
            for i in reversed(name_sanad_moein):
                final_name_sanad+=f'{i}'
            final_name_sanad+=f"-"
            for i in reversed(name_sanad_tafs):
                final_name_sanad+=f'{i}'

            ###################################
            row[3].text = f"{final_name_sanad}" 
            row[2].text = text_finaler 
            row[1].text = f'{obj_sanad.bedehkar}'
            row[0].text = f'{obj_sanad.bestankar}'



    print('1')

    # Adding data from the list to the table 


        # row = table.add_row().cells 
        # row[3].text = str(id) 
        # row[2].text = text_finaler 
        # row[1].text = "1200"
        # row[0].text = '2000'

    
    # Adding style to a table 
    table.style = 'Colorful List'
    
    print('1')

    # paragraph = doc.add_paragraph(style='rtl').add_run("سلام خوش اومدی")

    doc.add_paragraph(style='rtl').add_run("")

    ##############################################################################
    table2 = doc.add_table(rows=1, cols=4) 
    

    bedehkar_total = "{:,}".format(bedehkar_total)
    bestankar_total = "{:,}".format(bestankar_total)


    # Adding heading in the 1st row of the table 
    row = table2.rows[0].cells 
    row[3].text = 'بدهکار مجموع'
    row[2].text = 'بستانکار مجموع'

    print('1')
    
    # Adding data from the list to the table 

    
    row = table2.add_row().cells 
    row[3].text = f"{bedehkar_total}"
    row[2].text = f"{bestankar_total}"

    
    # Adding style to a table 
    table2.style = 'Colorful List'

    print('1')

    ###############################################################################


    # paragraph = doc.add_paragraph(style='rtl').add_run("سلام خوش اومدی")
    # # Increasing size of the font 
    # paragraph.font.size = Pt(12) 
    # paragraph.font.complex_script = True

    # para.font.rtl = True

    print('1')


    # Now save the document to a location 
    doc.save('gfg.docx') 
